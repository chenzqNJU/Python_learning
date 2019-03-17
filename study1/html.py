import pandas as pd
import numpy as np
from sgmllib import SGMLParser
import os
import sys
import urllib
import urllib2
from docx import Document
from docx.shared import Inches
import time
import json
from html.parser import HTMLParser

#####################################
import xlrd
import urllib.request


class MyHTMLParser(HTMLParser):  # 定义HTMLParser的子类,用以复写HTMLParser中的方法

    # 构造方法,定义data数组用来存储html中的数据
    def __init__(self):
        HTMLParser.__init__(self)
        self.data = []

    # 覆盖starttag方法,可以进行一些打印操作
    def handle_starttag(self, tag, attrs):
        pass
        # print("Start Tag: ",tag)
        # for attr in attrs:
        #   print(attr)

    # 覆盖endtag方法
    def handle_endtag(self, tag):
        pass

    # 覆盖handle_data方法,用来处理获取的html数据,这里保存在data数组
    def handle_data(self, data):
        if data.count('\n') == 0:
            self.data.append(data)


mainData_book = xlrd.open_workbook("E:\\stock\\公司公告.xls", formatting_info=True)
mainData_sheet = mainData_book.sheet_by_index(0)
row = 24
rowValues = mainData_sheet.row_values(row, start_colx=0, end_colx=4)
print('报告' in rowValues[2], rowValues[2])
url = rowValues[3].split('"')[1]

###############找出半年报
url = []
for row in range(20):
    rowValues = mainData_sheet.row_values(row, start_colx=0, end_colx=4)
    # print('报告' in rowValues[2],rowValues[2])
    if '报告' in rowValues[2]:
        print(rowValues[2])
        url.append(rowValues[3].split('"')[1])
url = url[1]

res = urllib.request.urlopen(url)
html = res.read().decode('utf-8')

parser = MyHTMLParser()  # 创建子类实例
parser.feed(html)  # 将html数据传给解析器进行解析
data = parser.data  # 对解析后的数据进行相应操作并打印


class MyHTMLParser(HTMLParser):

    def handle_starttag(self, tag, attrs):
        """
        recognize start tag, like <div>
        :param tag:
        :param attrs:
        :return:
        """
        print("Encountered a start tag:", tag)

    def handle_endtag(self, tag):
        """
        recognize end tag, like </div>
        :param tag:
        :return:
        """
        print("Encountered an end tag :", tag)

    # def handle_data(self, data):
    #     """
    #     recognize data, html content string
    #     :param data:
    #     :return:
    #     """
    #     print("Encountered some data  :", data)

    def handle_data(self, data):
        if self.lasttag == 'p':
            print("Encountered p data  :", data)

    def handle_startendtag(self, tag, attrs):
        """
        recognize tag that without endtag, like <img />
        :param tag:
        :param attrs:
        :return:
        """
        print("Encountered startendtag :", tag)

    def handle_comment(self, data):
        """

        :param data:
        :return:
        """
        print("Encountered comment :", data)


filenames = []
import re

for index, table in enumerate(res.xpath('//table')):
    caption = table.xpath(
        'string(./caption)').extract_first()  # 提取caption tag里面的所有text，包括子节点内的和文本子节点，这样也行 caption = ''.join(table.xpath('./caption//text()').extract())
    filename = str(index + 1) + '_' + caption if caption else str(index + 1)  # xpath 要用到 table 计数，从[1]开始
    filenames.append(re.sub(r'[^\w\s()]', '', filename))  # 移除特殊符号

parser = MyHTMLParser()
parser.feed('<html><head><title>Test</title></head>'
            '<body><h1>Parse me!</h1><img src = "" />'
            '<!-- comment --></body></html>')
parser.handle_data('<html><head><title>Test</title></head>'
                   '<body><h1>Parse me!</h1><img src = "" />'
                   '<!-- comment --></body></html>')
parser.data

# !/usr/bin/env python3
# _*_ coding=utf-8 _*_
##################################################
import csv
from urllib.request import urlopen
from bs4 import BeautifulSoup
from urllib.request import HTTPError

try:
    html = urlopen("http://en.wikipedia.org/wiki/Comparison_of_text_editors")
except HTTPError as e:
    print("not found")
bsObj = BeautifulSoup(html, "html.parser")
table = bsObj.findAll("table", {"class": "wikitable"})[0]
len(table)
if table is None:
    print("no table");
    exit(1)
rows = table.findAll("tr")

csvFile = open("editors.csv", 'wt', newline='', encoding='utf-8')
writer = csv.writer(csvFile)

################################################
###############找出报告
url = []
for row in range(20):
    rowValues = mainData_sheet.row_values(row, start_colx=0, end_colx=4)
    # print('报告' in rowValues[2],rowValues[2])
    if '报告' in rowValues[2]:
        print(rowValues[2])
        url.append(rowValues[3].split('"')[1])

table = bsObj.findAll("table")
len(table)
for t in range(len(table)):
    rows = table[t].findAll('tr')
    if '净利润' in str(rows): print(t)

rows = table[6].findAll('tr')
len(rows)

url1 = url[1]
url1 = 'http://snap.windin.com/ns/bulletin.php?code=BE146D53A9FF&id=95219842&type=1'

###############找出年报
mainData_book = xlrd.open_workbook("E:\\stock\\公司公告.xls", formatting_info=True)
mainData_sheet = mainData_book.sheet_by_index(0)
mainData_sheet.row_values(2, start_colx=1, end_colx=7)
len(mainData_sheet)
url = [];code = []
for row in range(103):
    rowValues = mainData_sheet.row_values(row, start_colx=0, end_colx=4)
    # print('报告' in rowValues[2],rowValues[2])
    if '年度报告' in rowValues[2] and rowValues[2].endswith('报告'):
        print(rowValues[2])
        code.append(rowValues[1])
        url.append(rowValues[3].split('"')[1])
for x in range(len(url)):
    res = urllib.request.urlopen(url[x])
    html = res.read().decode('utf-8')
    if len(html)>10000:print(x)
url[2]

res = urllib.request.urlopen('http://snap.windin.com/ns/getatt.php?id=95219850&att_id=38928467')
html = res.read().decode('gbk')
len(html)

res = urllib.request.urlopen('http://www.sse.com.cn/disclosure/listedinfo/announcement/c/2018-09-14/600509_2017_nB.pdf')

res = urllib.request.urlopen('http://data.eastmoney.com/notices/detail/000001/AN201808151178957195,JWU1JWI5JWIzJWU1JWFlJTg5JWU5JTkzJWI2JWU4JWExJThj.html')

html = res.read().decode('gbk')
' 证券事务代表' in html
n=html.find('其中：应付递延奖金')
html[n-100:n+799]

bsObj = BeautifulSoup(html, "html.parser")
table = bsObj.findAll("table")
len(table)

table[1]
for t in range(len(table)):
    rows = table[t].findAll('tr')
    if '主营业务分行业情况' in str(rows): print(t)
table[12]
rows = table[12].findAll('tr')
len(rows)

#########################################################
mat = None
for x, row in enumerate(rows):
    csvRow = []
    for cell in row.findAll(['td', 'th']):
        # cell = row.findAll(['td', 'th'])[0]
        if 'colspan' in str(cell):
            p1 = "(?<=colspan=\").+?(?=\")"
            n = int(re.search(re.compile(p1), str(cell)).group(0))
            csvRow = csvRow + [cell.get_text()] + [''] * (n - 1)
            continue
        csvRow.append(cell.get_text())
    # print(len(csvRow))
    # print(csvRow)
    if mat is None:
        mat = csvRow
    else:
        mat = np.vstack((mat, csvRow))
df = pd.DataFrame(mat)
###################################################

try:
    for row in rows:
        csvRow = []
        for cell in row.findAll(['td', 'th']):
            csvRow.append(cell.get_text())

        writer.writerow(csvRow)
finally:
    csvFile.close()

z = parser.feed(html)

from html.parser import HTMLParser
from html.entitydefs import name2codepoint


class MyHTMLParser(HTMLParser):

    def handle_starttag(self, tag, attrs):
        print('<%s>' % tag)

    def handle_endtag(self, tag):
        print('</%s>' % tag)

    def handle_startendtag(self, tag, attrs):
        print('<%s/>' % tag)

    def handle_data(self, data):
        print('data')

    def handle_comment(self, data):
        print('<!-- -->')

    def handle_entityref(self, name):
        print('&%s;' % name)

    def handle_charref(self, name):
        print('&#%s;' % name)


parser = MyHTMLParser()
parser.feed('<html><head></head><body><p>Some <a href=\"#\">html</a> tutorial...<br>END</p></body></html>')

#####################################
import json
# For python 3.x
from html.parser import HTMLParser


# 定义HTMLParser的子类,用以复写HTMLParser中的方法
class MyHTMLParser(HTMLParser):

    # 构造方法,定义data数组用来存储html中的数据
    def __init__(self):
        HTMLParser.__init__(self)
        self.data = []

    # 覆盖starttag方法,可以进行一些打印操作
    def handle_starttag(self, tag, attrs):
        pass
        # print("Start Tag: ",tag)
        # for attr in attrs:
        #   print(attr)

    # 覆盖endtag方法
    def handle_endtag(self, tag):
        pass

    # 覆盖handle_data方法,用来处理获取的html数据,这里保存在data数组
    def handle_data(self, data):
        if data.count('\n') == 0:
            self.data.append(data)


# 创建子类实例
parser = MyHTMLParser()

# 将html数据传给解析器进行解析
parser.feed(html)

# 对解析后的数据进行相应操作并打印
z = parser.data
parser.data[4]
for item in parser.data:
    if item.startswith("{\"columns\""):
        payloadDict = json.loads(item)
        list = payloadDict["payload"]["rows"]
        for backlog in list:
            if backlog[1] == "Product Backlog Item" or backlog[1] == "Bug":
                print(backlog[2], "       Point: ", backlog[3])

'''
##获取要解析的url
class GetUrl(SGMLParser):
    def __init__(self):
        SGMLParser.__init__(self)
        self.start = False
        self.urlArr = []

    def start_div(self, attr):
        for name, value in attr:
            if value == "ChairmanCont Bureau":  # 页面js中的固定值
                self.start = True

    def end_div(self):
        self.start = False

    def start_a(self, attr):
        if self.start:
            for name, value in attr:
                self.urlArr.append(value)

    def getUrlArr(self):
        return self.urlArr


##解析上面获取的url，获取有用数据
class getManInfo(SGMLParser):
    def __init__(self):
        SGMLParser.__init__(self)
        self.start = False
        self.p = False
        self.dl = False
        self.manInfo = []
        self.subInfo = []

    def start_div(self, attr):
        for name, value in attr:
            if value == "SpeakerInfo":  # 页面js中的固定值
                self.start = True

    def end_div(self):
        self.start = False

    def start_p(self, attr):
        if self.dl:
            self.p = True

    def end_p(self):
        self.p = False

    def start_img(self, attr):
        if self.dl:
            for name, value in attr:
                self.subInfo.append(value)

    def handle_data(self, data):
        if self.p:
            self.subInfo.append(data.decode('utf-8'))

    def start_dl(self, attr):
        if self.start:
            self.dl = True

    def end_dl(self):
        self.manInfo.append(self.subInfo)
        self.subInfo = []
        self.dl = False

    def getManInfo(self):
        return self.manInfo


urlSource = "http://www.XXX"
sourceData = urllib2.urlopen(urlSource).read()

startTime = time.clock()
##get urls
getUrl = GetUrl()
getUrl.feed(sourceData)
urlArr = getUrl.getUrlArr()
getUrl.close()
print
"get url use:" + str((time.clock() - startTime))
startTime = time.clock()

##get maninfos
manInfos = getManInfo()
for url in urlArr:  # one url one person
    data = urllib2.urlopen(url).read()
    manInfos.feed(data)
infos = manInfos.getManInfo()
manInfos.close()
print
"get maninfos use:" + str((time.clock() - startTime))
startTime = time.clock()

# word
saveFile = os.getcwd() + "\\xxx.docx"
doc = Document()
##word title
doc.add_heading("HEAD".decode('gbk'), 0)
p = doc.add_paragraph("HEADCONTENT:".decode('gbk'))

##write info
for infoArr in infos:
    i = 0
    for info in infoArr:
        if i == 0:  ##img url
            arr1 = info.split('.')
            suffix = arr1[len(arr1) - 1]
            arr2 = info.split('/')
            preffix = arr2[len(arr2) - 2]
            imgFile = os.getcwd() + "\\imgs\\" + preffix + "." + suffix
            if not os.path.exists(os.getcwd() + "\\imgs"):
                os.mkdir(os.getcwd() + "\\imgs")
            imgData = urllib2.urlopen(info).read()

            try:
                f = open(imgFile, 'wb')
                f.write(imgData)
                f.close()
                doc.add_picture(imgFile, width=Inches(1.25))
                os.remove(imgFile)
            except Exception as err:
                print(err)


        elif i == 1:
            doc.add_heading(info + ":", level=1)
        else:
            doc.add_paragraph(info, style='ListBullet')
        i = i + 1

doc.save(saveFile)
'''
